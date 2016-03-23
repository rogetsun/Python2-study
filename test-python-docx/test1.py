# coding:utf-8
__author__ = 'uv2sun'

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_TABLE_DIRECTION
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# Set a cell background (shading) color to RGB D9D9D9.
shading_elm = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
print shading_elm
recordset = [
    {'qty': 'q1', 'id': 'id1', 'desc': 'desc1'},
    {'qty': 'q2', 'id': 'id2', 'desc': 'desc2'},
    {'qty': 'q3', 'id': 'id3', 'desc': 'desc3'}
]

document = Document()
table1 = document.add_table(rows=1, cols=3)
# table.style = ("ColorfulList", "TableGrid")
table1.style = 'TableGrid'

table1.alignment = WD_TABLE_ALIGNMENT.CENTER
table1.direction = WD_TABLE_DIRECTION.LTR

hdr_cells = table1.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
for cs in hdr_cells:
    cs.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
for item in recordset:
    row_cells = table1.add_row().cells
    row_cells[0].text = str(item.get('qty'))
    row_cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    row_cells[0]._tc.get_or_add_tcPr().append(shading_elm)
    row_cells[1].text = str(item.get('id'))
    row_cells[2].text = item.get('desc')

# table2
table = document.add_table(rows=2, cols=3)
table.style = 'Colorful List Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = True
row1 = table.rows[0].cells
row1[0].merge(row1[2])

row1[0].text = 'Title'

hdr_cells = table.rows[1].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'

for item in recordset:
    row_cells = table.add_row().cells
    row_cells[0].text = str(item.get('qty'))
    row_cells[1].text = str(item.get('id'))
    row_cells[2].text = item.get('desc')

row_cells = table1.add_row().cells
row_cells[0].text = 'litx'
row_cells[1].text = 'litx2'
row_cells[2].text = 'litc3'

document.add_page_break()

document.save('demo.docx')
